import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime
import io
import base64
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import docx
from docx.shared import Inches, Pt, RGBColor
from pptx import Presentation
from pptx.util import Inches, Pt
import streamlit as st

class ReportGenerator:
    def __init__(self):
        self.styles = getSampleStyleSheet()
    
    def create_monthly_report(self, month, defect_stats, issue_analysis, shipment_data, complaint_data):
        """生成客诉月报 - B.3"""
        
        # 创建数据摘要
        report_summary = {
            '报告月份': month,
            '生成时间': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            '总客诉数': len(complaint_data) if complaint_data is not None else 0,
            '涉及机型数': complaint_data['机型_标准化'].nunique() if complaint_data is not None and '机型_标准化' in complaint_data.columns else 0,
            '总出货数': shipment_data['出货数'].sum() if shipment_data is not None and '出货数' in shipment_data.columns else 0,
            '平均不良率': defect_stats['不良率(%)'].mean() if defect_stats is not None and not defect_stats.empty else 0,
            '主要问题类型': issue_analysis.index[0] if issue_analysis is not None and not issue_analysis.empty else '无'
        }
        
        return report_summary
    
    def create_visualizations(self, defect_stats, issue_analysis):
        """创建可视化图表"""
        figures = {}
        
        # 1. 不良率柱状图
        if defect_stats is not None and not defect_stats.empty:
            fig1 = px.bar(
                defect_stats, 
                x='机型_标准化', 
                y='不良率(%)',
                title='各机型不良率统计',
                color='不良率(%)',
                color_continuous_scale='Viridis'
            )
            fig1.update_layout(xaxis_title='机型', yaxis_title='不良率(%)')
            figures['defect_rate'] = fig1
        
        # 2. 问题分类饼图
        if issue_analysis is not None and not issue_analysis.empty:
            fig2 = px.pie(
                issue_analysis.reset_index(),
                values='问题数量',
                names='问题分类',
                title='客诉问题分类分布',
                hole=0.3
            )
            figures['issue_distribution'] = fig2
        
        # 3. 不良数趋势图
        if defect_stats is not None and not defect_stats.empty:
            fig3 = px.line(
                defect_stats,
                x='机型_标准化',
                y='不良数',
                title='各机型不良数',
                markers=True
            )
            fig3.update_traces(line=dict(width=3))
            figures['defect_trend'] = fig3
        
        return figures
    
    def export_to_word(self, report_summary, defect_stats, issue_analysis, filename="客诉分析报告.docx"):
        """导出Word报告"""
        doc = docx.Document()
        
        # 标题
        title = doc.add_heading('客诉数据分析报告', 0)
        title.alignment = 1  # 居中
        
        # 报告摘要
        doc.add_heading('报告摘要', level=1)
        
        summary_table = doc.add_table(rows=len(report_summary)+1, cols=2)
        summary_table.style = 'Light Shading Accent 1'
        
        # 表头
        header_cells = summary_table.rows[0].cells
        header_cells[0].text = '指标'
        header_cells[1].text = '数值'
        
        # 数据行
        for i, (key, value) in enumerate(report_summary.items(), 1):
            row_cells = summary_table.rows[i].cells
            row_cells[0].text = str(key)
            row_cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # 不良率统计
        if defect_stats is not None and not defect_stats.empty:
            doc.add_heading('不良率统计', level=1)
            
            defect_table = doc.add_table(rows=len(defect_stats)+1, cols=len(defect_stats.columns))
            defect_table.style = 'Light Grid Accent 1'
            
            # 表头
            header_cells = defect_table.rows[0].cells
            for j, col in enumerate(defect_stats.columns):
                header_cells[j].text = str(col)
            
            # 数据行
            for i, row in defect_stats.iterrows():
                row_cells = defect_table.rows[i+1].cells
                for j, value in enumerate(row):
                    row_cells[j].text = str(value)
        
        # 保存文件
        doc.save(filename)
        return filename
    
    def export_to_pdf(self, report_summary, filename="客诉分析报告.pdf"):
        """导出PDF报告 (简化版)"""
        doc = SimpleDocTemplate(filename, pagesize=letter)
        elements = []
        
        # 标题
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            alignment=1,
            spaceAfter=30
        )
        elements.append(Paragraph('客诉数据分析报告', title_style))
        
        # 报告摘要
        elements.append(Paragraph('报告摘要', self.styles['Heading2']))
        elements.append(Spacer(1, 12))
        
        # 创建摘要表格
        summary_data = [['指标', '数值']]
        for key, value in report_summary.items():
            summary_data.append([str(key), str(value)])
        
        summary_table = Table(summary_data, colWidths=[3*inch, 3*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ]))
        
        elements.append(summary_table)
        elements.append(Spacer(1, 20))
        
        doc.build(elements)
        return filename
